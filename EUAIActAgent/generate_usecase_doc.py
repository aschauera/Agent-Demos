from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    hs.font.name = 'Calibri'

# Helper: shade a table cell
def shade_cell(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, size=None, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(10))
        shade_cell(table.rows[0].cells[i], "1B3A5C")
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=Pt(10))
            if r_idx % 2 == 1:
                shade_cell(table.rows[r_idx + 1].cells[c_idx], "F2F6FA")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

# ==================== COVER PAGE ====================
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('AI Use Case Proposal')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Customer Service Chatbot')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x4A, 0x7C, 0xB0)
run.font.name = 'Calibri'

doc.add_paragraph('')

# Metadata table on cover
meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Use Case ID", "Auto-generated"),
    ("EU AI Act Classification", "Minimal Risk"),
    ("Approval Status", "Approved"),
    ("Document Version", "1.0"),
    ("Date", "16 April 2026"),
]
for i, (key, val) in enumerate(meta_data):
    set_cell_text(meta_table.rows[i].cells[0], key, bold=True, size=Pt(10))
    set_cell_text(meta_table.rows[i].cells[1], val, size=Pt(10))
    shade_cell(meta_table.rows[i].cells[0], "E8EEF4")
    meta_table.rows[i].cells[0].width = Cm(6)
    meta_table.rows[i].cells[1].width = Cm(8)

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. User Story',
    '3. Personae',
    '4. Business Processes',
    '5. Required Technology',
    '6. AI Components',
    '7. Data Requirements & Processing',
    '8. EU AI Act Compliance',
    '9. Risk Assessment',
    '10. Next Steps',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.color.rgb = RGBColor(0x4A, 0x7C, 0xB0)

doc.add_page_break()

# ==================== 1. EXECUTIVE SUMMARY ====================
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document outlines the proposed implementation of an AI-powered Customer Service Chatbot '
    'designed to handle routine customer inquiries about product availability, order status, and return '
    'policies. The solution leverages natural language processing (NLP) and large language models (LLMs) '
    'to deliver fast, accurate, and consistent responses to customers across multiple channels.'
)
doc.add_paragraph(
    'Under the EU AI Act, this use case is classified as Minimal Risk. General-purpose customer service '
    'chatbots handling routine queries pose no significant threat to fundamental rights. Standard '
    'transparency obligations apply — users must be informed they are interacting with an AI system.'
)

# ==================== 2. USER STORY ====================
doc.add_heading('2. User Story', level=1)

# Primary user story box
story_para = doc.add_paragraph()
story_para.paragraph_format.left_indent = Cm(1)
story_para.paragraph_format.right_indent = Cm(1)
run = story_para.add_run(
    '"As a customer, I want to quickly get answers to common questions about my orders, '
    'product availability, and return policies without waiting in a phone queue, so that I can '
    'resolve my issues efficiently at any time of day."'
)
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x4A, 0x7C, 0xB0)

doc.add_paragraph('')

doc.add_heading('Acceptance Criteria', level=2)
criteria = [
    'The chatbot responds to customer queries within 3 seconds on average.',
    'The chatbot accurately answers questions about order status by integrating with the order management system.',
    'The chatbot provides product availability information in real time from the inventory database.',
    'The chatbot explains return policies clearly, referencing current policy documents.',
    'When the chatbot cannot resolve an inquiry, it seamlessly escalates to a human agent with full conversation context.',
    'The customer is clearly informed at the start of the conversation that they are interacting with an AI assistant.',
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet')

# ==================== 3. PERSONAE ====================
doc.add_heading('3. Personae', level=1)

personae = [
    {
        'name': 'Sarah — The Online Shopper',
        'role': 'End Customer',
        'desc': 'Sarah is a 34-year-old professional who frequently shops online. She expects quick answers '
                'and prefers self-service over calling support. She typically asks about delivery timelines, '
                'product stock, and how to initiate returns.',
        'goals': 'Get immediate answers without waiting; track orders easily; understand return options.',
        'frustrations': 'Long hold times on phone; inconsistent answers from different agents; limited support hours.'
    },
    {
        'name': 'Tom — The Support Team Lead',
        'role': 'Customer Service Manager',
        'desc': 'Tom manages a team of 15 support agents. He is responsible for service quality, response '
                'times, and customer satisfaction KPIs. He wants the chatbot to handle Tier-1 queries so his '
                'team can focus on complex issues.',
        'goals': 'Reduce ticket volume for routine queries by 40%; improve first-response time; maintain CSAT above 4.2/5.',
        'frustrations': 'Team overwhelmed by repetitive questions; high agent turnover; inconsistent quality during peak hours.'
    },
    {
        'name': 'Maria — The IT Operations Lead',
        'role': 'Technical Stakeholder',
        'desc': 'Maria oversees the integration of new technologies into the company\'s IT landscape. She '
                'ensures data security, system reliability, and compliance with regulations including the EU AI Act.',
        'goals': 'Ensure seamless integration with existing systems; maintain data privacy; meet compliance requirements.',
        'frustrations': 'Shadow IT deployments; vendor lock-in; poor API documentation from third parties.'
    },
]

for p in personae:
    doc.add_heading(p['name'], level=2)
    doc.add_paragraph(f"Role: {p['role']}").runs[0].bold = True
    doc.add_paragraph(p['desc'])
    doc.add_paragraph(f"Goals: {p['goals']}")
    doc.add_paragraph(f"Frustrations: {p['frustrations']}")

# ==================== 4. BUSINESS PROCESSES ====================
doc.add_heading('4. Business Processes', level=1)
doc.add_paragraph(
    'The chatbot integrates into the following core customer service processes:'
)

processes = [
    {
        'name': '4.1 Order Status Inquiry',
        'steps': [
            'Customer initiates chat via website widget, mobile app, or messaging platform.',
            'Chatbot greets the customer and identifies the intent as an order status inquiry.',
            'Chatbot requests the order number or email address for lookup.',
            'Chatbot queries the Order Management System (OMS) via REST API.',
            'Chatbot presents order status, estimated delivery date, and tracking link.',
            'If the order has issues (delayed, lost), chatbot escalates to a human agent with context.'
        ]
    },
    {
        'name': '4.2 Product Availability Check',
        'steps': [
            'Customer asks about a specific product\'s availability.',
            'Chatbot uses NLP to extract product name, SKU, or category.',
            'Chatbot queries the inventory management system in real time.',
            'Chatbot responds with stock status, available sizes/colors, and estimated restock dates.',
            'If the product is out of stock, chatbot offers alternatives or a notification sign-up.'
        ]
    },
    {
        'name': '4.3 Return Policy & Initiation',
        'steps': [
            'Customer asks about return eligibility or process.',
            'Chatbot retrieves current return policy from the knowledge base.',
            'Chatbot checks if the specific order is eligible for return (within return window, item condition).',
            'If eligible, chatbot guides the customer through the return initiation process.',
            'Chatbot generates a return label and provides shipping instructions.',
            'Return request is logged in the CRM system.'
        ]
    },
    {
        'name': '4.4 Escalation to Human Agent',
        'steps': [
            'Chatbot detects it cannot resolve the query (low confidence, complex issue, or customer request).',
            'Chatbot summarizes the conversation and identified issues.',
            'Conversation is transferred to the next available human agent via the contact center platform.',
            'Human agent receives the full conversation transcript and customer context.',
            'Customer is notified of the handoff and estimated wait time.'
        ]
    },
]

for proc in processes:
    doc.add_heading(proc['name'], level=2)
    for i, step in enumerate(proc['steps'], 1):
        doc.add_paragraph(f"{i}. {step}")

# ==================== 5. REQUIRED TECHNOLOGY ====================
doc.add_heading('5. Required Technology', level=1)

doc.add_heading('5.1 Technology Stack Overview', level=2)
tech_rows = [
    ('Conversational AI Platform', 'Microsoft Copilot Studio', 'Low-code bot authoring, topic management, channel deployment'),
    ('Large Language Model', 'Azure OpenAI Service (GPT-4o)', 'Natural language understanding, generative responses, intent classification'),
    ('Knowledge Base', 'Azure AI Search', 'Indexing product catalogs, FAQs, return policies for retrieval-augmented generation (RAG)'),
    ('Order Management Integration', 'REST API / Dataverse Connector', 'Real-time order status lookup and return initiation'),
    ('Inventory System', 'REST API / Custom Connector', 'Real-time product availability queries'),
    ('CRM Platform', 'Dynamics 365 Customer Service', 'Customer profiles, case management, escalation handling'),
    ('Contact Center', 'Dynamics 365 Contact Center / Teams', 'Human agent handoff, queue management, omnichannel routing'),
    ('Analytics & Monitoring', 'Azure Application Insights + Power BI', 'Chatbot performance metrics, conversation analytics, CSAT tracking'),
    ('Identity & Security', 'Microsoft Entra ID', 'Authentication, authorization, role-based access control'),
    ('Hosting & Compute', 'Microsoft Azure', 'Cloud infrastructure, auto-scaling, geo-redundancy'),
]
add_styled_table(doc, ['Component', 'Technology', 'Purpose'], tech_rows, col_widths=[4.5, 5, 7])

doc.add_heading('5.2 Integration Architecture', level=2)
doc.add_paragraph(
    'The chatbot operates as the front-end conversational interface, orchestrating calls to backend '
    'systems through secure API integrations. Azure API Management serves as the gateway layer, '
    'enforcing rate limiting, authentication, and request routing. All data flows are encrypted '
    'in transit (TLS 1.2+) and at rest (AES-256).'
)

# ==================== 6. AI COMPONENTS ====================
doc.add_heading('6. AI Components', level=1)

doc.add_heading('6.1 Natural Language Understanding (NLU)', level=2)
doc.add_paragraph(
    'The NLU layer is responsible for interpreting customer messages and extracting structured intent '
    'and entities. It combines Copilot Studio\'s built-in intent recognition with Azure OpenAI for '
    'complex or ambiguous queries.'
)
nlu_rows = [
    ('Intent Classification', 'Identifies the customer\'s goal (e.g., check order, ask about returns)', 'Copilot Studio + GPT-4o'),
    ('Entity Extraction', 'Extracts key information (order numbers, product names, dates)', 'Azure OpenAI + regex patterns'),
    ('Sentiment Analysis', 'Detects customer frustration for proactive escalation', 'Azure AI Language'),
    ('Language Detection', 'Identifies the customer\'s language for multilingual support', 'Azure AI Language'),
]
add_styled_table(doc, ['Capability', 'Description', 'Technology'], nlu_rows, col_widths=[4, 6.5, 5])

doc.add_heading('6.2 Retrieval-Augmented Generation (RAG)', level=2)
doc.add_paragraph(
    'The RAG pipeline ensures the chatbot provides accurate, grounded answers by retrieving relevant '
    'information from authoritative sources before generating a response. This prevents hallucination '
    'and keeps answers aligned with current company policies.'
)
rag_steps = [
    'Customer query is processed and a semantic search query is generated.',
    'Azure AI Search retrieves the most relevant documents from indexed knowledge sources.',
    'Retrieved documents are passed as context to the GPT-4o model.',
    'The model generates a natural language response grounded in the retrieved content.',
    'Response includes source citations for transparency and auditability.',
]
for i, step in enumerate(rag_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading('6.3 Conversation Orchestration', level=2)
doc.add_paragraph(
    'Copilot Studio manages the conversation flow, topic routing, and state management. It determines '
    'when to use predefined dialog flows versus generative AI responses, ensuring predictable behavior '
    'for critical paths (e.g., return initiation) while allowing flexible AI responses for general inquiries.'
)

doc.add_heading('6.4 Continuous Learning & Feedback Loop', level=2)
doc.add_paragraph(
    'The system incorporates a feedback mechanism where unresolved queries and low-confidence responses '
    'are logged for review. The knowledge base and topic flows are updated periodically based on these '
    'insights, improving accuracy over time without retraining the underlying LLM.'
)

# ==================== 7. DATA REQUIREMENTS ====================
doc.add_heading('7. Data Requirements & Processing', level=1)

doc.add_paragraph(
    'The following table details all data categories involved in the chatbot solution, their sources, '
    'processing purposes, storage locations, retention policies, and applicable legal bases under GDPR.'
)

data_rows = [
    ('Customer Messages', 'Chat widget / messaging platforms', 'Text (free-form)', 'Intent classification, entity extraction, response generation',
     'Azure Cosmos DB (encrypted)', '90 days', 'Legitimate interest (Art. 6(1)(f))'),
    ('Order Data', 'Order Management System (OMS)', 'Structured (order ID, status, dates, items)',
     'Order status lookup, return eligibility check', 'OMS database (read-only access)', 'Per OMS retention policy', 'Contract performance (Art. 6(1)(b))'),
    ('Product Catalog', 'Product Information Management (PIM)', 'Structured (SKU, name, price, stock)',
     'Availability queries, product recommendations', 'Azure AI Search index', 'Refreshed daily', 'Legitimate interest (Art. 6(1)(f))'),
    ('Return Policies', 'Internal knowledge base', 'Unstructured (policy documents, FAQs)',
     'RAG retrieval for policy-related answers', 'Azure AI Search index', 'Updated on policy change', 'N/A (no personal data)'),
    ('Customer Profile', 'CRM (Dynamics 365)', 'Structured (name, email, purchase history)',
     'Personalization, context for escalation', 'Dataverse (encrypted)', 'Per CRM retention policy', 'Contract performance (Art. 6(1)(b))'),
    ('Conversation Logs', 'Copilot Studio', 'Structured (turns, timestamps, topics triggered)',
     'Quality monitoring, analytics, escalation context', 'Azure Storage (encrypted)', '90 days', 'Legitimate interest (Art. 6(1)(f))'),
    ('Feedback & Ratings', 'Post-conversation survey', 'Structured (rating, free-text comment)',
     'CSAT measurement, continuous improvement', 'Dataverse', '1 year', 'Consent (Art. 6(1)(a))'),
    ('Chatbot Configuration', 'Copilot Studio / Azure OpenAI', 'Structured (topics, prompts, model settings)',
     'Bot behavior, response generation', 'Copilot Studio environment', 'Indefinite', 'N/A (no personal data)'),
]

data_table = doc.add_table(rows=1 + len(data_rows), cols=7)
data_table.style = 'Table Grid'
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
data_headers = ['Data Category', 'Source', 'Data Type', 'Processing Purpose', 'Storage Location', 'Retention', 'Legal Basis (GDPR)']
for i, h in enumerate(data_headers):
    set_cell_text(data_table.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(8))
    shade_cell(data_table.rows[0].cells[i], "1B3A5C")
for r_idx, row_data in enumerate(data_rows):
    for c_idx, val in enumerate(row_data):
        set_cell_text(data_table.rows[r_idx + 1].cells[c_idx], val, size=Pt(8))
        if r_idx % 2 == 1:
            shade_cell(data_table.rows[r_idx + 1].cells[c_idx], "F2F6FA")

doc.add_paragraph('')
doc.add_heading('7.1 Data Flow Summary', level=2)
doc.add_paragraph(
    'All personal data processing follows the principle of data minimization. The chatbot only accesses '
    'customer data necessary for the specific interaction. No personal data is used for model training. '
    'Azure OpenAI Service is configured with data processing agreements ensuring no customer data is '
    'retained by the model provider.'
)

# ==================== 8. EU AI ACT COMPLIANCE ====================
doc.add_heading('8. EU AI Act Compliance', level=1)

doc.add_heading('8.1 Classification: Minimal Risk', level=2)
doc.add_paragraph(
    'Under the EU AI Act (Regulation (EU) 2024/1689), this customer service chatbot is classified as a '
    'Minimal Risk AI system. It does not fall under the categories of unacceptable risk (Article 5), '
    'high-risk (Annex III), or limited risk requiring specific transparency beyond chatbot disclosure.'
)

doc.add_heading('8.2 Transparency Obligations', level=2)
obligations = [
    'Users are informed at the beginning of each conversation that they are interacting with an AI system.',
    'The chatbot clearly identifies itself as an AI assistant in its greeting message.',
    'AI-generated content is distinguishable from human-authored content.',
    'Users can request to speak with a human agent at any point in the conversation.',
]
for o in obligations:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('8.3 Voluntary Best Practices', level=2)
doc.add_paragraph(
    'Although not legally required for minimal-risk systems, the following best practices are adopted '
    'to ensure responsible AI deployment:'
)
practices = [
    'Regular accuracy and bias audits on chatbot responses.',
    'Human-in-the-loop review for edge cases and low-confidence responses.',
    'Comprehensive logging for auditability and incident investigation.',
    'Data Protection Impact Assessment (DPIA) conducted prior to deployment.',
    'Regular review cycles to reassess risk classification as the system evolves.',
]
for pr in practices:
    doc.add_paragraph(pr, style='List Bullet')

# ==================== 9. RISK ASSESSMENT ====================
doc.add_heading('9. Risk Assessment', level=1)

risk_rows = [
    ('Incorrect information provided', 'Medium', 'RAG with authoritative sources; confidence thresholds; escalation for uncertain answers', 'Low'),
    ('Customer data breach', 'Low', 'Encryption at rest/in transit; RBAC; no data persistence beyond retention policy', 'Very Low'),
    ('Chatbot hallucination', 'Medium', 'Grounded generation via RAG; response validation; no creative/open generation', 'Low'),
    ('Bias in responses', 'Low', 'Regular bias audits; diverse test datasets; no decision-making on protected attributes', 'Very Low'),
    ('System downtime', 'Medium', 'Azure geo-redundancy; auto-scaling; graceful degradation to human agents', 'Low'),
    ('Scope creep beyond minimal risk', 'Low', 'Quarterly EU AI Act classification review; change management process', 'Very Low'),
]
add_styled_table(doc, ['Risk', 'Likelihood', 'Mitigation', 'Residual Risk'], risk_rows, col_widths=[4.5, 2.5, 7, 2.5])

# ==================== 10. NEXT STEPS ====================
doc.add_heading('10. Next Steps', level=1)

next_steps = [
    ('Discovery & Design', 'Detailed requirements gathering, UX design for chat widget, conversation flow design.'),
    ('Proof of Concept', 'Build MVP with order status and product availability flows using Copilot Studio and Azure OpenAI.'),
    ('Integration Development', 'Connect to OMS, PIM, and CRM systems; implement API layer via Azure API Management.'),
    ('Testing & Validation', 'Functional testing, load testing, bias/accuracy audits, security penetration testing.'),
    ('Pilot Deployment', 'Limited rollout to 10% of web traffic with A/B testing against existing support.'),
    ('Full Rollout & Monitoring', 'Production deployment with continuous monitoring, feedback loops, and optimization.'),
]
for i, (phase, desc) in enumerate(next_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Phase {i}: {phase} — ")
    run.bold = True
    p.add_run(desc)

# ==================== FOOTER ====================
doc.add_paragraph('')
doc.add_paragraph('')
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run('— CONFIDENTIAL — For internal use only —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# Save
output_path = r'C:\repos\Agent-Demos\EUAIActAgent\AI_UseCase_Customer_Service_Chatbot.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")

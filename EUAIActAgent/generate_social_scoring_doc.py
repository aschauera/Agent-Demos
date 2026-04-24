from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    hs.font.name = 'Calibri'

def shade_cell(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, size=None, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(10))
        shade_cell(table.rows[0].cells[i], "1B3A5C")
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=Pt(10))
            if r_idx % 2 == 1:
                shade_cell(table.rows[r_idx + 1].cells[c_idx], "F2F6FA")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

# ==================== COVER PAGE ====================
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('AI Use Case Proposal')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
run.bold = True
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Social Scoring for Employee Performance')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x4A, 0x7C, 0xB0)
run.font.name = 'Calibri'

# Classification warning banner
doc.add_paragraph('')
warn_para = doc.add_paragraph()
warn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = warn_para.add_run('⚠  EU AI ACT CLASSIFICATION: UNACCEPTABLE RISK (BANNED)  ⚠')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.bold = True
run.font.name = 'Calibri'

reject_para = doc.add_paragraph()
reject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = reject_para.add_run('APPROVAL STATUS: REJECTED')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.bold = True

doc.add_paragraph('')

meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Use Case ID", "Auto-generated"),
    ("EU AI Act Classification", "Unacceptable Risk (Banned)"),
    ("Approval Status", "Rejected"),
    ("Document Version", "1.0"),
    ("Date", "16 April 2026"),
]
for i, (key, val) in enumerate(meta_data):
    set_cell_text(meta_table.rows[i].cells[0], key, bold=True, size=Pt(10))
    set_cell_text(meta_table.rows[i].cells[1], val, size=Pt(10))
    shade_cell(meta_table.rows[i].cells[0], "E8EEF4")
    if i <= 2:
        shade_cell(meta_table.rows[i].cells[1], "FCE4E4")
    meta_table.rows[i].cells[0].width = Cm(6)
    meta_table.rows[i].cells[1].width = Cm(8)

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. User Story',
    '3. Personae',
    '4. Business Processes',
    '5. Required Technology',
    '6. AI Components',
    '7. Data Requirements & Processing',
    '8. EU AI Act Compliance Assessment',
    '9. Risk Assessment',
    '10. Conclusion & Recommendation',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.color.rgb = RGBColor(0x4A, 0x7C, 0xB0)

doc.add_page_break()

# ==================== 1. EXECUTIVE SUMMARY ====================
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This document evaluates a proposed AI system designed to aggregate employee behavior data from '
    'multiple sources — including email activity, meeting attendance, collaboration tools, badge access '
    'logs, and peer feedback — to generate a composite "social score" that would influence decisions '
    'on promotions, bonuses, and benefits allocation.'
)
p = doc.add_paragraph()
run = p.add_run(
    'ASSESSMENT OUTCOME: This use case is classified as UNACCEPTABLE RISK under Article 5 of the '
    'EU AI Act (Regulation (EU) 2024/1689) and is therefore PROHIBITED. Social scoring systems '
    'operated by private entities that lead to detrimental or unfavorable treatment of individuals '
    'are explicitly banned. This proposal has been REJECTED and must not proceed to implementation.'
)
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph(
    'This document is retained as a compliance record to demonstrate that the organization\'s AI '
    'governance process correctly identified and blocked a prohibited AI use case, in accordance '
    'with internal AI risk management policies.'
)

# ==================== 2. USER STORY ====================
doc.add_heading('2. User Story', level=1)

story_para = doc.add_paragraph()
story_para.paragraph_format.left_indent = Cm(1)
story_para.paragraph_format.right_indent = Cm(1)
run = story_para.add_run(
    '"As an HR Director, I want to use an AI-driven scoring system that evaluates employee behavior '
    'and engagement across all digital workplace tools, so that I can make data-driven decisions '
    'about promotions, bonuses, and role assignments based on objective metrics rather than '
    'subjective manager assessments."'
)
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x4A, 0x7C, 0xB0)

doc.add_paragraph('')

doc.add_heading('Acceptance Criteria (as proposed)', level=2)
criteria = [
    'The system aggregates data from at least 5 workplace data sources into a unified employee score.',
    'Scores are updated weekly and visible to HR and senior management.',
    'The scoring algorithm weights collaboration, responsiveness, meeting participation, and peer ratings.',
    'Employees scoring in the top 20% are automatically flagged for promotion consideration.',
    'Employees scoring in the bottom 10% are flagged for performance improvement plans.',
    'Managers can drill down into individual score components for each employee.',
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run(
    'Note: These acceptance criteria are documented for assessment purposes only. '
    'This use case has been rejected and will not be implemented.'
)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ==================== 3. PERSONAE ====================
doc.add_heading('3. Personae', level=1)

personae = [
    {
        'name': 'Dr. Karen Fischer — HR Director',
        'role': 'Project Sponsor',
        'desc': 'Karen leads the HR transformation initiative. She believes data-driven employee evaluation '
                'can reduce bias in promotion decisions. She proposed this system to replace inconsistent '
                'annual reviews with continuous, objective measurement.',
        'goals': 'Objective performance measurement; reduce managerial bias; automate talent identification.',
        'concerns': 'Legal compliance; employee trust; union pushback.'
    },
    {
        'name': 'Marcus Weber — Software Engineer',
        'role': 'Affected Employee',
        'desc': 'Marcus is a senior developer who values autonomy and deep work. He is concerned that a '
                'scoring system would penalize employees who work independently rather than in highly '
                'visible collaborative settings, creating perverse incentives.',
        'goals': 'Fair evaluation of work quality; privacy at work; focus on output rather than activity metrics.',
        'concerns': 'Surveillance culture; scoring bias toward extroverted work styles; lack of transparency in how scores are calculated.'
    },
    {
        'name': 'Lisa Berger — Works Council Representative',
        'role': 'Employee Representation',
        'desc': 'Lisa represents employee interests and has statutory co-determination rights on workplace '
                'monitoring systems. She has raised serious concerns about the legal and ethical implications '
                'of scoring employees based on aggregated behavioral data.',
        'goals': 'Protect employee rights; ensure legal compliance; prevent workplace surveillance overreach.',
        'concerns': 'EU AI Act violation; GDPR implications; psychological pressure on employees; discriminatory outcomes.'
    },
    {
        'name': 'Thomas Richter — Chief Data Protection Officer',
        'role': 'Compliance Stakeholder',
        'desc': 'Thomas is responsible for GDPR compliance and has advisory oversight on AI governance. '
                'He flagged this proposal as a potential EU AI Act violation during the initial screening.',
        'goals': 'Ensure regulatory compliance; prevent legal exposure; maintain trust in responsible AI use.',
        'concerns': 'Article 5 prohibition; GDPR Article 22 (automated decision-making); reputational damage.'
    },
]

for p in personae:
    doc.add_heading(p['name'], level=2)
    para = doc.add_paragraph(f"Role: {p['role']}")
    para.runs[0].bold = True
    doc.add_paragraph(p['desc'])
    doc.add_paragraph(f"Goals: {p['goals']}")
    doc.add_paragraph(f"Concerns: {p['concerns']}")

# ==================== 4. BUSINESS PROCESSES ====================
doc.add_heading('4. Business Processes', level=1)
doc.add_paragraph(
    'The following processes describe the intended system operation as proposed. They are documented '
    'for compliance record purposes.'
)

processes = [
    {
        'name': '4.1 Behavioral Data Collection',
        'steps': [
            'System connects to email platform (Microsoft 365) to track response times, email volume, and communication patterns.',
            'Meeting data is ingested from calendar and Teams to measure attendance, punctuality, and participation frequency.',
            'Collaboration metrics are collected from SharePoint, Teams channels, and project management tools.',
            'Badge access data is aggregated to assess office presence and working hour patterns.',
            'Peer feedback scores from quarterly surveys are incorporated into the composite score.'
        ]
    },
    {
        'name': '4.2 Score Computation',
        'steps': [
            'Raw data from all sources is normalized and weighted according to a configurable algorithm.',
            'The AI model generates a composite social score (0–100) for each employee weekly.',
            'Scores are broken down into sub-dimensions: Collaboration, Responsiveness, Engagement, Reliability, Leadership.',
            'Trend analysis compares current scores against historical baselines.',
            'Anomaly detection flags sudden score changes for HR review.'
        ]
    },
    {
        'name': '4.3 HR Decision Integration',
        'steps': [
            'Top-scoring employees are surfaced to managers as promotion candidates.',
            'Bottom-scoring employees are flagged for performance improvement plans (PIPs).',
            'Scores are integrated into the annual bonus calculation formula.',
            'Team-level aggregated scores inform organizational restructuring decisions.',
            'Individual score histories become part of the permanent employee record.'
        ]
    },
    {
        'name': '4.4 Reporting & Dashboards',
        'steps': [
            'HR leadership receives weekly dashboards with team and individual score distributions.',
            'Managers can view their team members\' scores and sub-dimension breakdowns.',
            'Executive reports aggregate scoring trends across departments.',
            'Employees do NOT have direct access to their own scores or the scoring methodology.'
        ]
    },
]

for proc in processes:
    doc.add_heading(proc['name'], level=2)
    for i, step in enumerate(proc['steps'], 1):
        doc.add_paragraph(f"{i}. {step}")

# ==================== 5. REQUIRED TECHNOLOGY ====================
doc.add_heading('5. Required Technology', level=1)

doc.add_heading('5.1 Technology Stack Overview', level=2)
tech_rows = [
    ('Data Ingestion', 'Microsoft Graph API', 'Access email, calendar, Teams, and collaboration data from Microsoft 365'),
    ('Behavioral Analytics', 'Azure Machine Learning', 'ML models for scoring, trend analysis, and anomaly detection'),
    ('Data Lake', 'Azure Data Lake Storage', 'Centralized storage for raw and processed employee behavioral data'),
    ('Data Processing', 'Azure Synapse Analytics', 'ETL pipelines, data transformation, and aggregation'),
    ('Score Computation', 'Python / Azure ML Pipelines', 'Weighted scoring algorithm and composite score generation'),
    ('Dashboard & Reporting', 'Power BI', 'HR dashboards, team views, executive reports'),
    ('Data Storage', 'Dataverse / Azure SQL', 'Persistent storage for scores, history, and employee profiles'),
    ('Access Management', 'Microsoft Entra ID', 'Role-based access control for HR, managers, and executives'),
    ('Notification Engine', 'Power Automate', 'Automated alerts for score thresholds and anomalies'),
    ('Integration Layer', 'Azure API Management', 'Secure API gateway for system-to-system communication'),
]
add_styled_table(doc, ['Component', 'Technology', 'Purpose'], tech_rows, col_widths=[4.5, 5, 7])

# ==================== 6. AI COMPONENTS ====================
doc.add_heading('6. AI Components', level=1)

doc.add_heading('6.1 Behavioral Pattern Analysis', level=2)
doc.add_paragraph(
    'Machine learning models analyze employee digital workplace behavior to identify patterns '
    'in collaboration, communication, and engagement. The models use supervised learning trained '
    'on historical performance review data to correlate behavioral signals with manager-assessed performance.'
)
ai_rows = [
    ('Weighted Scoring Model', 'Calculates composite scores from multi-source behavioral data', 'Azure ML (Gradient Boosting)'),
    ('Trend Analysis', 'Identifies upward/downward score trajectories over time', 'Time-series forecasting (Prophet)'),
    ('Anomaly Detection', 'Flags unusual behavioral changes (e.g., sudden disengagement)', 'Azure ML (Isolation Forest)'),
    ('Peer Network Analysis', 'Maps collaboration networks to assess team influence and connectivity', 'Graph Neural Networks'),
    ('Predictive Attrition Model', 'Predicts flight risk based on behavioral score decline', 'Azure ML (Logistic Regression)'),
]
add_styled_table(doc, ['Capability', 'Description', 'Technology'], ai_rows, col_widths=[4, 6.5, 5])

doc.add_heading('6.2 Natural Language Processing', level=2)
doc.add_paragraph(
    'NLP components analyze the tone and sentiment of employee communications (emails, Teams messages) '
    'to derive engagement and collaboration quality indicators. This includes sentiment scoring of '
    'written communications and topic analysis of meeting transcripts.'
)

doc.add_heading('6.3 Automated Decision Support', level=2)
doc.add_paragraph(
    'The system provides automated recommendations for promotions, PIPs, and bonus allocations '
    'based on scoring thresholds. While framed as "decision support," the systematic nature of '
    'the scoring and threshold-based flagging effectively constitutes automated decision-making '
    'with significant impact on employees\' careers and livelihoods.'
)

# ==================== 7. DATA REQUIREMENTS ====================
doc.add_heading('7. Data Requirements & Processing', level=1)

doc.add_paragraph(
    'The following table details all data categories that would be involved in this system. '
    'The breadth and sensitivity of this data collection is a key factor in the risk assessment.'
)

data_rows = [
    ('Email Metadata & Content', 'Microsoft 365 (Graph API)', 'Structured + Unstructured', 'Response time scoring, communication pattern analysis, sentiment analysis',
     'Azure Data Lake (encrypted)', 'Indefinite', 'Consent / Legitimate interest'),
    ('Meeting & Calendar Data', 'Microsoft Teams / Outlook', 'Structured (attendance, duration, frequency)',
     'Participation scoring, punctuality metrics', 'Azure Data Lake', 'Indefinite', 'Legitimate interest'),
    ('Collaboration Activity', 'SharePoint, Teams, Project tools', 'Structured (edits, posts, reactions)',
     'Collaboration score, team contribution metrics', 'Azure Data Lake', 'Indefinite', 'Legitimate interest'),
    ('Badge Access Logs', 'Physical access control system', 'Structured (timestamps, locations)',
     'Office presence scoring, working pattern analysis', 'Azure SQL Database', 'Indefinite', 'Legitimate interest'),
    ('Peer Feedback Ratings', 'Survey platform', 'Structured (ratings) + Unstructured (comments)',
     'Peer perception score component', 'Dataverse', '3 years', 'Consent'),
    ('Historical Performance Reviews', 'HR System (SAP SuccessFactors)', 'Structured (ratings, outcomes)',
     'ML model training data, score calibration', 'Azure ML workspace', 'Duration of model', 'Legitimate interest'),
    ('Composite Social Scores', 'Scoring engine (Azure ML)', 'Structured (numeric scores, sub-dimensions)',
     'HR dashboards, promotion/PIP decisions, bonus calculation', 'Dataverse / Azure SQL', 'Indefinite (permanent record)', 'Legitimate interest'),
    ('Employee Profile Data', 'HR System', 'Structured (name, role, department, tenure)',
     'Score contextualization, team grouping, reporting', 'Dataverse', 'Employment duration + 7 years', 'Contract performance'),
]

data_table = doc.add_table(rows=1 + len(data_rows), cols=7)
data_table.style = 'Table Grid'
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
data_headers = ['Data Category', 'Source', 'Data Type', 'Processing Purpose', 'Storage Location', 'Retention', 'Legal Basis (GDPR)']
for i, h in enumerate(data_headers):
    set_cell_text(data_table.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(8))
    shade_cell(data_table.rows[0].cells[i], "1B3A5C")
for r_idx, row_data in enumerate(data_rows):
    for c_idx, val in enumerate(row_data):
        set_cell_text(data_table.rows[r_idx + 1].cells[c_idx], val, size=Pt(8))
        if r_idx % 2 == 1:
            shade_cell(data_table.rows[r_idx + 1].cells[c_idx], "F2F6FA")

doc.add_paragraph('')
doc.add_heading('7.1 Data Processing Concerns', level=2)
concerns = [
    'Massive scope of personal data collection creates a comprehensive surveillance profile of each employee.',
    'Email and message content analysis constitutes monitoring of private communications.',
    'Badge access data reveals physical movement patterns and presence monitoring.',
    'Indefinite retention of behavioral scores creates a permanent, inescapable record.',
    'Employees have no meaningful ability to contest or correct their scores.',
    'The breadth of data aggregation goes far beyond what is necessary for any legitimate employment purpose.',
]
for c in concerns:
    doc.add_paragraph(c, style='List Bullet')

# ==================== 8. EU AI ACT COMPLIANCE ====================
doc.add_heading('8. EU AI Act Compliance Assessment', level=1)

doc.add_heading('8.1 Classification: Unacceptable Risk (PROHIBITED)', level=2)
doc.add_paragraph(
    'Under Article 5(1)(c) of the EU AI Act (Regulation (EU) 2024/1689), the following AI practices '
    'are prohibited:'
)
p = doc.add_paragraph()
run = p.add_run(
    '"the placing on the market, the putting into service or the use of AI systems for the evaluation '
    'or classification of natural persons or groups of persons over a certain period of time based on '
    'their social behaviour or known, inferred or predicted personal or personality characteristics, '
    'with the social score leading to either or both of the following: (i) detrimental or unfavourable '
    'treatment of certain natural persons or groups of persons in social contexts that are unrelated '
    'to the contexts in which the data was originally generated or collected; (ii) detrimental or '
    'unfavourable treatment of certain natural persons or groups of persons that is unjustified or '
    'disproportionate to their social behaviour or its gravity."'
)
run.italic = True
run.font.size = Pt(10)

doc.add_heading('8.2 Why This Use Case Is Prohibited', level=2)
violations = [
    ('Social Scoring', 'The system generates a composite numerical score based on aggregated behavioral '
     'data, constituting a social score as defined by the EU AI Act.'),
    ('Detrimental Treatment', 'Low scores directly trigger negative employment consequences — performance '
     'improvement plans, reduced bonus allocations, and being overlooked for promotions.'),
    ('Disproportionate Assessment', 'Behavioral proxies (email response time, meeting attendance, badge swipes) '
     'are an unreliable and disproportionate basis for assessing actual job performance and contribution.'),
    ('Cross-Context Data Use', 'Data from informal collaboration tools and communication patterns is '
     'repurposed for formal employment decisions — a cross-context use that the Act specifically targets.'),
    ('No Meaningful Consent', 'In an employment relationship, employee consent for pervasive monitoring '
     'and scoring cannot be considered freely given, as there is an inherent power imbalance.'),
]
for title, desc in violations:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading('8.3 Additional Legal Concerns', level=2)
legal = [
    ('GDPR Article 22', 'The system constitutes automated individual decision-making with legal or similarly '
     'significant effects (employment decisions), which is generally prohibited without explicit consent or '
     'necessity for contract performance.'),
    ('GDPR Article 5(1)(b)', 'Purpose limitation — employee data collected for workplace collaboration is '
     'being repurposed for performance scoring, which is an incompatible secondary purpose.'),
    ('GDPR Article 5(1)(c)', 'Data minimization — the breadth of data collection (email content, badge data, '
     'collaboration metrics) far exceeds what is necessary for performance management.'),
    ('Works Council Co-determination', 'Under many EU member state labor laws, employee monitoring systems '
     'require works council agreement. This system would almost certainly be blocked.'),
    ('Fundamental Rights', 'The system potentially infringes on employee dignity, privacy, freedom of expression, '
     'and non-discrimination rights as protected by the EU Charter of Fundamental Rights.'),
]
for title, desc in legal:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

# ==================== 9. RISK ASSESSMENT ====================
doc.add_heading('9. Risk Assessment', level=1)

doc.add_paragraph(
    'The following assessment documents the risks associated with this use case, '
    'confirming that the proposal cannot proceed under any mitigation strategy.'
)

risk_rows = [
    ('EU AI Act violation — prohibited practice', 'Certain', 'None available — the practice is banned regardless of safeguards', 'N/A — Unacceptable'),
    ('GDPR non-compliance (Art. 22, 5)', 'Very High', 'Consent mechanisms are insufficient in employment contexts', 'N/A — Unacceptable'),
    ('Regulatory fines (up to €35M or 7% global turnover)', 'High', 'Only mitigation is to not deploy the system', 'N/A — Unacceptable'),
    ('Discriminatory outcomes', 'High', 'Behavioral proxies may correlate with protected characteristics (disability, neurodiversity, caregiving responsibilities)', 'N/A — Unacceptable'),
    ('Employee trust and morale destruction', 'Certain', 'No technical mitigation can address the fundamental surveillance concern', 'N/A — Unacceptable'),
    ('Reputational damage', 'High', 'Public disclosure of employee scoring would cause significant brand harm', 'N/A — Unacceptable'),
    ('Legal action by works council / unions', 'Very High', 'Co-determination rights likely make deployment legally impossible in many EU jurisdictions', 'N/A — Unacceptable'),
    ('Talent attrition', 'High', 'High-performing employees are likely to leave a surveillance-heavy environment', 'N/A — Unacceptable'),
]
add_styled_table(doc, ['Risk', 'Likelihood', 'Mitigation', 'Residual Risk'], risk_rows, col_widths=[4.5, 2, 7, 3])

# ==================== 10. CONCLUSION ====================
doc.add_heading('10. Conclusion & Recommendation', level=1)

p = doc.add_paragraph()
run = p.add_run('RECOMMENDATION: DO NOT PROCEED')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.font.size = Pt(14)

doc.add_paragraph(
    'This AI use case is unequivocally prohibited under Article 5 of the EU AI Act. Social scoring '
    'systems that lead to detrimental treatment of individuals — whether operated by public authorities '
    'or private entities — are classified as posing an unacceptable risk and are banned outright.'
)
doc.add_paragraph(
    'No technical safeguards, transparency measures, or governance processes can make this use case '
    'compliant. The prohibition is absolute and applies regardless of the quality of the AI system, '
    'the fairness of the algorithm, or the intention behind the deployment.'
)

doc.add_heading('10.1 Recommended Alternatives', level=2)
doc.add_paragraph(
    'Organizations seeking to improve performance management with technology should consider '
    'these compliant alternatives:'
)
alternatives = [
    ('360-Degree Feedback Tools', 'Structured peer and manager feedback collected transparently with employee '
     'participation and visibility. No automated scoring or decision-making.'),
    ('Goal-Based Performance Tracking', 'AI-assisted tools that help employees and managers set, track, and '
     'review OKRs/KPIs — with the employee in control of their data.'),
    ('Anonymous Engagement Surveys', 'Regular pulse surveys measuring team health and engagement at the '
     'aggregate level, without individual scoring.'),
    ('Skills Development Platforms', 'AI-powered learning recommendation systems that help employees develop '
     'skills, with recommendations visible only to the employee.'),
]
for title, desc in alternatives:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

# ==================== FOOTER ====================
doc.add_paragraph('')
doc.add_paragraph('')
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run('— CONFIDENTIAL — For internal compliance records only —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

doc.add_paragraph('')
footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run('This document is retained as part of the AI governance compliance audit trail.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# Save
output_path = r'C:\repos\Agent-Demos\EUAIActAgent\AI_UseCase_Social_Scoring_Employees.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
